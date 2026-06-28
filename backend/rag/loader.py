from langchain_community.document_loaders import PyPDFLoader, TextLoader
from langchain.schema import Document
from pathlib import Path
import csv
import logging

logger = logging.getLogger(__name__)


def _load_csv(file_path: str, name: str) -> list[Document]:
    """Read a CSV into a single text Document (header-aware)."""
    rows = []
    with open(file_path, newline="", encoding="utf-8", errors="ignore") as f:
        reader = csv.reader(f)
        for row in reader:
            rows.append(", ".join(cell.strip() for cell in row))
    text = "\n".join(rows)
    logger.info(f"Loaded CSV: {name} — {len(rows)} rows")
    return [Document(page_content=text, metadata={"source": name})]


def _load_docx(file_path: str, name: str) -> list[Document]:
    """Extract text from a .docx using python-docx."""
    from docx import Document as DocxDocument  # python-docx

    doc = DocxDocument(file_path)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    # Include table cell text too.
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    text = "\n".join(parts)
    logger.info(f"Loaded DOCX: {name} — {len(parts)} blocks")
    return [Document(page_content=text, metadata={"source": name})]


def load_document(file_path: str) -> list[Document]:
    """
    Load a document based on file extension.
    Supports: PDF, TXT, MD, DOCX, CSV
    """
    path = Path(file_path)
    extension = path.suffix.lower()

    try:
        if extension == ".pdf":
            loader = PyPDFLoader(file_path)
            documents = loader.load()
            logger.info(f"Loaded PDF: {path.name} — {len(documents)} pages")
            return documents

        elif extension in [".txt", ".md"]:
            loader = TextLoader(file_path, encoding="utf-8")
            documents = loader.load()
            logger.info(f"Loaded text file: {path.name}")
            return documents

        elif extension == ".csv":
            return _load_csv(file_path, path.name)

        elif extension == ".docx":
            return _load_docx(file_path, path.name)

        else:
            raise ValueError(f"Unsupported file type: {extension}")

    except Exception as e:
        logger.error(f"Failed to load document {file_path}: {e}")
        raise
